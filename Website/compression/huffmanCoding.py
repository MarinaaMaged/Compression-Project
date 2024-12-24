import heapq
import os
from collections import Counter
from docx import Document  # For handling Word files
from PyPDF2 import PdfReader  # For handling PDF files


class HuffmanCoding:
    def __init__(self):
        self.heap = []
        self.codes = {}
        self.reverse_mapping = {}
        self.root = None

    class Node:
        def __init__(self, char, freq):
            self.char = char
            self.freq = freq
            self.left = None
            self.right = None

        def __lt__(self, other):
            return self.freq < other.freq

    def make_frequency_dict(self, text):
        return Counter(text)

    def build_heap(self, frequency):
        for key, freq in frequency.items():
            node = self.Node(key, freq)
            heapq.heappush(self.heap, node)

    def merge_nodes(self):
        while len(self.heap) > 1:
            node1 = heapq.heappop(self.heap)
            node2 = heapq.heappop(self.heap)

            merged = self.Node(None, node1.freq + node2.freq)
            merged.left = node1
            merged.right = node2
            heapq.heappush(self.heap, merged)

        self.root = self.heap[0] if self.heap else None

    def make_codes_helper(self, root, current_code):
        if root is None:
            return
        if root.char is not None:
            self.codes[root.char] = current_code
            self.reverse_mapping[current_code] = root.char
        self.make_codes_helper(root.left, current_code + "0")
        self.make_codes_helper(root.right, current_code + "1")

    def make_codes(self):
        self.make_codes_helper(self.root, "")

    def get_encoded_text(self, text):
        return "".join(self.codes[char] for char in text)

    def pad_encoded_text(self, encoded_text):
        extra_padding = 8 - len(encoded_text) % 8
        encoded_text += "0" * extra_padding
        return f"{extra_padding:08b}" + encoded_text

    def serialize_tree(self, root):
        if root is None:
            return "0"
        if root.char is not None:
            return f"1{root.char}"
        return f"2{self.serialize_tree(root.left)}{self.serialize_tree(root.right)}"

    def read_text_from_file(self, file_path):
        extension = os.path.splitext(file_path)[1].lower()
        if extension == ".txt":
            with open(file_path, 'r') as file:
                return file.read()
        elif extension == ".docx":
            document = Document(file_path)
            return "\n".join([p.text for p in document.paragraphs])
        elif extension == ".pdf":
            reader = PdfReader(file_path)
            return "".join(page.extract_text() for page in reader.pages)
        else:
            raise ValueError("Unsupported file type. Supported: .txt, .docx, .pdf")


    def compress(self, file_path):
        base_name, extension = os.path.splitext(file_path)
        output_path = f"{base_name}_compressed.txt"

        # Read file content based on file type
        if extension == ".txt":
            with open(file_path, 'r', encoding="utf-8") as file:
                text = file.read()
        elif extension == ".docx":
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif extension == ".pdf":
            reader = PdfReader(file_path)
            text = "\n".join([page.extract_text() for page in reader.pages])
        else:
            raise ValueError(f"Unsupported file type: {extension}")

        if not text.strip():
            raise ValueError("Input file is empty or contains only whitespace.")

        frequency = self.make_frequency_dict(text)
        self.build_heap(frequency)
        self.merge_nodes()
        self.make_codes()

        # Serialize tree and encode text
        serialized_tree = self.serialize_tree(self.root)
        encoded_text = self.get_encoded_text(text)
        padded_encoded_text = self.pad_encoded_text(encoded_text)

        # Save compressed data to a text file for clarity
        with open(output_path, 'w', encoding="utf-8") as output:
            tree_length = len(serialized_tree)
            output.write(f"{tree_length}\n")
            output.write(serialized_tree + "\n")
            output.write(padded_encoded_text)

        return output_path

    def decompress(self, file_path):
        base_name, extension = os.path.splitext(file_path)
        output_path = f"{base_name}_decompressed{extension}"

        with open(file_path, 'r', encoding="utf-8") as file:
            tree_length = int(file.readline().strip())
            serialized_tree = file.read(tree_length)
            bit_string = file.read().strip()

        root, _ = self.deserialize_tree(serialized_tree)
        self.root = root
        self.make_codes_helper(root, "")

        # Remove padding
        padded_info = bit_string[:8]
        extra_padding = int(padded_info, 2)
        bit_string = bit_string[8:-extra_padding]

        # Decode bit string to original text
        current_code = ""
        decoded_text = []
        for bit in bit_string:
            current_code += bit
            if current_code in self.reverse_mapping:
                decoded_text.append(self.reverse_mapping[current_code])
                current_code = ""

        decoded_text = "".join(decoded_text)

        # Write decompressed content back to appropriate file type
        if extension == ".txt":
            with open(output_path, 'w', encoding="utf-8") as output:
                output.write(decoded_text)
        elif extension == ".docx":
            doc = Document()
            for line in decoded_text.split("\n"):
                doc.add_paragraph(line)
            doc.save(output_path)
        elif extension == ".pdf":
            raise NotImplementedError("Decompression to PDF is not supported yet.")
        else:
            raise ValueError(f"Unsupported file type: {extension}")

        return output_path


    def deserialize_tree(self, serialized_tree):
        if not serialized_tree:
            return None, serialized_tree

        node_type = serialized_tree[0]
        serialized_tree = serialized_tree[1:]

        if node_type == "0":
            return None, serialized_tree
        elif node_type == "1":
            char = serialized_tree[0]
            return self.Node(char, 0), serialized_tree[1:]
        elif node_type == "2":
            left, serialized_tree = self.deserialize_tree(serialized_tree)
            right, serialized_tree = self.deserialize_tree(serialized_tree)
            node = self.Node(None, 0)
            node.left = left
            node.right = right
            return node, serialized_tree
